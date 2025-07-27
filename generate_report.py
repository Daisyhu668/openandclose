import os
import re
from datetime import datetime
from typing import List, Dict

try:
    from PIL import Image
    import pytesseract
    from docx import Document
    from docx.shared import Inches
except ImportError as exc:
    raise SystemExit("Required packages are missing. Please install Pillow, pytesseract and python-docx") from exc


BASE_DIR = os.path.dirname(os.path.abspath(__file__))
TEMPLATE_NAME = "抵押物走访模版.docx"

# -------------- helpers -----------------

def find_certificate(files: List[str]) -> str | None:
    """Return path of the certificate image based on filename."""
    for f in files:
        name = os.path.basename(f)
        if any(keyword in name for keyword in ("房产证", "不动产权证")):
            return f
    return None


def ocr_image(path: str) -> str:
    """Run OCR on the given image and return extracted text."""
    img = Image.open(path)
    return pytesseract.image_to_string(img, lang="chi_sim")


def parse_certificate_text(text: str) -> Dict[str, str]:
    fields = {
        "权利人": r"权利人[:：]?([^\n]+)",
        "坐落": r"坐落[:：]?([^\n]+)",
        "建筑面积": r"建筑面积[:：]?([0-9\.]+[^\n]*)",
    }
    data: Dict[str, str] = {}
    for key, pattern in fields.items():
        m = re.search(pattern, text)
        data[key] = m.group(1).strip() if m else "未识别"
    return data


def classify_photos(files: List[str]) -> Dict[str, List[str]]:
    info: Dict[str, List[str]] = {
        "走访照片": [],
        "贝壳查询": [],
        "安居客查询": [],
        "评估": []
    }
    for f in files:
        name = os.path.basename(f)
        if name.startswith("走访"):
            info["走访照片"].append(f)
        elif "贝壳" in name:
            info["贝壳查询"].append(f)
        elif "安居客" in name:
            info["安居客查询"].append(f)
        elif "评估" in name or "查询" in name:
            info["评估"].append(f)
    return info


def replace_placeholder(doc: Document, placeholder: str, value: str) -> None:
    for p in doc.paragraphs:
        if placeholder in p.text:
            p.text = p.text.replace(placeholder, value)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                replace_placeholder(cell, placeholder, value)


def insert_photo(doc: Document, placeholder: str, path: str) -> None:
    for p in doc.paragraphs:
        if placeholder in p.text:
            p.text = ""
            run = p.add_run()
            run.add_picture(path, width=Inches(2.5))
            return


def main() -> None:
    files = [os.path.join(BASE_DIR, f) for f in os.listdir(BASE_DIR) if f.lower().endswith((".png", ".jpg", ".jpeg"))]
    cert_path = find_certificate(files)
    if not cert_path:
        raise SystemExit("房产证图片未找到")

    cert_text = ocr_image(cert_path)
    cert_data = parse_certificate_text(cert_text)
    owner = cert_data.get("权利人", "未知客户")

    other_files = [f for f in files if f != cert_path]
    photos = classify_photos(other_files)

    template_path = os.path.join(BASE_DIR, TEMPLATE_NAME)
    if not os.path.exists(template_path):
        raise SystemExit(f"模板文件 {TEMPLATE_NAME} 不存在")

    doc = Document(template_path)

    for key, value in cert_data.items():
        placeholder = f"【{key}】"
        replace_placeholder(doc, placeholder, value)
    replace_placeholder(doc, "【调查日期】", datetime.now().strftime("%Y年%m月%d日"))

    placeholders = [
        "【请在此处插入照片1：写字楼外观】",
        "【请在此处插入照片2：写字楼大门】",
        "【请在此处插入照片3：抵押物内景】",
        "【请在此处插入照片4：抵押物内景】",
        "【请在此处插入照片5：抵押物门牌】",
        "【请在此处插入照片6：抵押物消防设施】",
    ]

    for idx, img_path in enumerate(photos["走访照片"]):
        if idx < len(placeholders):
            insert_photo(doc, placeholders[idx], img_path)

    output_path = os.path.join(BASE_DIR, f"{owner}-抵押物走访报告.docx")
    doc.save(output_path)
    print("报告生成于", output_path)


if __name__ == "__main__":
    main()
