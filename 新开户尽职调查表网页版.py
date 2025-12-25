#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""新开户尽职调查表 · 智能生成脚本

功能亮点：
- 固定模板：`新开户尽职调查表模版.docx`。
- 仅替换三个占位符：{{客户名称}} / {{行业分类}} / {{经营地址}}。
- “上门核实图片”段落下自动插入照片，每页正好 2 张，并在纵横混合时自动做白底等比填充。
- 若模板中缺失占位符，会立即报错提醒，避免生成空白字段。
- 支持 CLI 与交互模式，可配置默认照片目录、PDF 导出等。
"""

from __future__ import annotations

import argparse
import json
import os
import re
import shutil
import subprocess
import sys
from typing import Dict, Iterable, List

from docx import Document
from docx.enum.table import WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.shared import Inches

try:
    from PIL import Image, ImageOps

    PIL_OK = True
except Exception:  # pragma: no cover - pillow 可选
    PIL_OK = False


# ===== 路径与常量 =====
CURRENT_DIR = os.path.dirname(os.path.abspath(__file__))
TEMPLATE_PATH_FIXED = os.path.join(CURRENT_DIR, "新开户尽职调查表模版.docx")
OUT_DIR_FIXED = CURRENT_DIR
DEFAULT_PHOTOS_DIR = os.path.join(CURRENT_DIR, "photos")
CONFIG_PATH = os.path.expanduser("~/.auto_dd_config.json")
REPORT_SUFFIX = "新开户尽职调查表"

PHOTO_TOKENS = ("{{走访照片}}", "{{照片}}", "{{照片区}}")
PHOTO_SECTION_KEYWORDS = ("上门核实图片",)
PLACEHOLDER_KEYS = ("客户名称", "行业分类", "经营地址")

CANVAS_SIZE = (2200, 1650)  # letterbox 4:3
CANVAS_MARGIN = 80
PHOTOS_PER_PAGE = 2
PHOTO_WIDTH_IN = 5.6
PHOTO_ROW_HEIGHT_IN = 4.6


# ===== 配置 =====
def load_config() -> Dict[str, str]:
    cfg = {
        "template": TEMPLATE_PATH_FIXED,
        "photos": DEFAULT_PHOTOS_DIR,
        "out": OUT_DIR_FIXED,
    }
    if os.path.exists(CONFIG_PATH):
        try:
            with open(CONFIG_PATH, "r", encoding="utf-8") as fh:
                data = json.load(fh)
            if isinstance(data, dict):
                cfg.update({k: v for k, v in data.items() if v})
        except Exception:
            pass
    cfg["template"] = os.environ.get("DUE_DILIGENCE_TEMPLATE", cfg["template"])
    cfg["photos"] = os.environ.get("DUE_DILIGENCE_PHOTOS", cfg["photos"])
    cfg["out"] = os.environ.get("DUE_DILIGENCE_OUT", cfg["out"])
    return cfg


def save_config(cfg: Dict[str, str]) -> None:
    try:
        with open(CONFIG_PATH, "w", encoding="utf-8") as fh:
            json.dump(cfg, fh, ensure_ascii=False, indent=2)
    except Exception:
        pass


# ===== 基础工具 =====
def expand_photos_arg(arg: str, debug: bool = False, recursive: bool = True) -> List[str]:
    if not arg:
        return []
    path = arg.strip()
    exts = (".jpg", ".jpeg", ".png", ".bmp", ".gif", ".webp", ".tif", ".tiff", ".heic", ".heif")
    out: List[str] = []

    if os.path.isdir(path):
        walker: Iterable = os.walk(path) if recursive else [(path, [], os.listdir(path))]
        for root, _, files in walker:
            for name in files:
                if name.lower().endswith(exts):
                    full = os.path.join(root, name)
                    out.append(full)
                    if debug:
                        print("  📷", full)
        out.sort()
        if debug:
            print(f"🔎 目录 {path} → {len(out)} 张")
        return out

    for piece in (x.strip() for x in path.split(",")):
        if not piece:
            continue
        if os.path.exists(piece):
            out.append(piece)
            if debug:
                print("  📷", piece)
        elif debug:
            print("  ❌ 未找到：", piece)
    if debug:
        print(f"🔎 文件列表 → {len(out)} 张")
    return out


def next_nonconflicting_path(path: str) -> str:
    if not os.path.exists(path):
        return path
    root, ext = os.path.splitext(path)
    idx = 1
    while True:
        candidate = f"{root} ({idx}){ext}"
        if not os.path.exists(candidate):
            return candidate
        idx += 1


def sanitize_filename(name: str) -> str:
    safe = (name or "未命名").strip() or "未命名"
    return re.sub(r"[\\/:*?\"<>|]", "_", safe)


# ===== 占位符替换 =====
def _apply_patterns(text: str, patterns: Dict[str, re.Pattern], values: Dict[str, str]):
    new_text = text
    changed = False
    for key, pattern in patterns.items():
        new_text, count = pattern.subn(values[key], new_text)
        if count:
            changed = True
    return new_text, changed


def replace_placeholders(doc: Document, values: Dict[str, str]) -> None:
    patterns = {
        key: re.compile(r"\{\{\s*" + re.escape(key) + r"\s*\}\}")
        for key in values
    }

    def process(paragraph) -> None:
        for run in paragraph.runs:
            new_text, changed = _apply_patterns(run.text, patterns, values)
            if changed:
                run.text = new_text
        combined = paragraph.text or ""
        new_combined, changed = _apply_patterns(combined, patterns, values)
        if changed and new_combined != combined:
            paragraph.text = new_combined

    for para in doc.paragraphs:
        process(para)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    process(para)


def inspect_template_placeholders(path: str, keys: Iterable[str]) -> Dict[str, int]:
    if not os.path.exists(path):
        raise FileNotFoundError(f"模板不存在：{path}")
    doc = Document(path)
    patterns = {
        key: re.compile(r"\{\{\s*" + re.escape(key) + r"\s*\}\}")
        for key in keys
    }
    stats = {key: 0 for key in keys}
    for para in _iter_paragraphs(doc):
        text = para.text or ""
        for key, pattern in patterns.items():
            stats[key] += len(pattern.findall(text))
    return stats


# ===== 照片处理 =====
def _letterbox_image(img, size: tuple[int, int], margin: int):
    canvas = Image.new("RGB", size, "white")
    target_w = max(size[0] - margin * 2, 1)
    target_h = max(size[1] - margin * 2, 1)
    scale = min(target_w / img.width, target_h / img.height, 1.0)
    resized = img.resize((int(img.width * scale), int(img.height * scale)), Image.LANCZOS)
    offset = ((size[0] - resized.width) // 2, (size[1] - resized.height) // 2)
    canvas.paste(resized, offset)
    return canvas


def _prepare_photo_for_word(src: str, cache_dir: str, target_px: int = 1600, debug: bool = False) -> str:
    os.makedirs(cache_dir, exist_ok=True)
    base, _ = os.path.splitext(os.path.basename(src))
    out_jpg = os.path.join(cache_dir, f"{base}_processed.jpg")

    if src.lower().endswith((".heic", ".heif")):
        sips = shutil.which("sips")
        if sips:
            try:
                subprocess.run([sips, "-s", "format", "jpeg", src, "--out", out_jpg], check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
                if os.path.exists(out_jpg) and os.path.getsize(out_jpg) > 1024:
                    return out_jpg
            except Exception as exc:
                if debug:
                    print("  ❌ HEIC 转换失败：", exc)
        return ""

    if not PIL_OK:
        return src

    try:
        with Image.open(src) as img:
            img = ImageOps.exif_transpose(img)
            largest = max(img.size)
            if largest > target_px:
                scale = target_px / largest
                img = img.resize((int(img.width * scale), int(img.height * scale)), Image.LANCZOS)
            if img.mode != "RGB":
                img = img.convert("RGB")
            boxed = _letterbox_image(img, CANVAS_SIZE, CANVAS_MARGIN)
            boxed.save(out_jpg, "JPEG", quality=88, optimize=True)
            return out_jpg
    except Exception as exc:
        if debug:
            print("  ❌ 照片处理失败：", exc)
        try:
            shutil.copy2(src, out_jpg)
            return out_jpg
        except Exception:
            return ""


def _iter_paragraphs(doc: Document):
    for para in doc.paragraphs:
        yield para
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    yield para


def _clear_photo_tokens(paragraph) -> None:
    for token in PHOTO_TOKENS:
        if token in paragraph.text:
            paragraph.text = paragraph.text.replace(token, "")


def _find_photo_anchor(doc: Document):
    for para in _iter_paragraphs(doc):
        if any(token in (para.text or "") for token in PHOTO_TOKENS):
            return para
    for para in _iter_paragraphs(doc):
        txt = (para.text or "").strip()
        if txt and any(keyword in txt for keyword in PHOTO_SECTION_KEYWORDS):
            return para
    return doc.add_paragraph("上门核实图片：")


def insert_photos(doc: Document, photos: List[str], out_dir: str, debug: bool = False) -> None:
    if not photos:
        if debug:
            print("ℹ️ 无照片需要插入。")
        return

    cache_dir = os.path.join(out_dir, "_photo_cache")
    processed: List[str] = []
    for path in photos:
        if not os.path.exists(path):
            if debug:
                print("  ❌ 缺少照片：", path)
            continue
        cooked = _prepare_photo_for_word(path, cache_dir, target_px=1800, debug=debug)
        if cooked:
            processed.append(cooked)
    if not processed:
        if debug:
            print("❌ 无可用照片，跳过插入。")
        return

    anchor = _find_photo_anchor(doc)
    _clear_photo_tokens(anchor)
    last_element = anchor._p

    groups = [processed[i : i + PHOTOS_PER_PAGE] for i in range(0, len(processed), PHOTOS_PER_PAGE)]
    for idx, group in enumerate(groups):
        table = doc.add_table(rows=len(group), cols=1)
        try:
            table.style = "Table Grid"
        except Exception:
            pass
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for r, image_path in enumerate(group):
            row = table.rows[r]
            row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
            row.height = Inches(PHOTO_ROW_HEIGHT_IN)
            cell = row.cells[0]
            cell.text = ""
            if not cell.paragraphs:
                cell.add_paragraph()
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            try:
                run = para.add_run()
                run.add_picture(image_path, width=Inches(PHOTO_WIDTH_IN))
            except Exception as exc:
                para.text = f"[图片插入失败：{os.path.basename(image_path)}]\n{exc}"
        last_element.addnext(table._tbl)
        last_element = table._tbl
        if idx < len(groups) - 1:
            breaker = doc.add_paragraph()
            breaker.add_run().add_break(WD_BREAK.PAGE)
            last_element.addnext(breaker._p)
            last_element = breaker._p

    if debug:
        print(f"  ✅ 已插入 {len(processed)} 张照片（每页 2 张）。")


# ===== 生成流程 =====
def generate_report(
    fields: Dict[str, str],
    photos: List[str],
    want_pdf: bool = False,
    open_out: bool = False,
    reveal_out: bool = False,
    debug: bool = False,
) -> str:
    cfg = load_config()
    template_path = cfg.get("template", TEMPLATE_PATH_FIXED)
    outdir = cfg.get("out", OUT_DIR_FIXED)
    os.makedirs(outdir, exist_ok=True)

    if not os.path.exists(template_path):
        raise FileNotFoundError(f"模板不存在：{template_path}")

    placeholders = {key: fields.get(key, "") for key in PLACEHOLDER_KEYS}
    stats = inspect_template_placeholders(template_path, PLACEHOLDER_KEYS)
    missing = [key for key, count in stats.items() if count == 0]
    if missing:
        raise ValueError(
            "模板未找到占位符：" + ",".join(missing) + "。请确认模板内存在 {{字段}}。"
        )

    print("占位符命中：", "，".join(f"{k}:{stats[k]}" for k in PLACEHOLDER_KEYS))

    doc = Document(template_path)
    replace_placeholders(doc, placeholders)

    insert_photos(doc, photos, outdir, debug=debug)

    base_name = f"{sanitize_filename(fields.get('客户名称', '未命名'))}{REPORT_SUFFIX}.docx"
    docx_out = next_nonconflicting_path(os.path.join(outdir, base_name))
    doc.save(docx_out)
    print("✅ 已生成 DOCX：", os.path.abspath(docx_out))

    if want_pdf:
        pdf_out = docx_out.replace(".docx", ".pdf")
        ok, tool = try_pdf(docx_out, pdf_out)
        if ok:
            print(f"✅ 已导出 PDF（{tool}）：{os.path.abspath(pdf_out)}")
        else:
            print("⚠️ PDF 转换失败（未检测到 Word/libreoffice）。")

    if reveal_out:
        try:
            subprocess.run(["open", "-R", docx_out])
        except Exception:
            pass
    elif open_out:
        try:
            subprocess.run(["open", os.path.dirname(docx_out)])
        except Exception:
            pass

    print("—— 完成 ——")
    return docx_out


# ===== PDF 支持 =====
def try_pdf(docx_path: str, pdf_path: str):
    try:
        from docx2pdf import convert as d2p

        d2p(docx_path, pdf_path)
        if os.path.exists(pdf_path) and os.path.getsize(pdf_path) > 1024:
            return True, "docx2pdf"
    except Exception:
        pass

    soffice = shutil.which("soffice") or "/Applications/LibreOffice.app/Contents/MacOS/soffice"
    if soffice and os.path.exists(soffice):
        try:
            outdir = os.path.dirname(docx_path)
            subprocess.run(
                [soffice, "--headless", "--convert-to", "pdf", "--outdir", outdir, docx_path],
                check=True,
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
            )
            if os.path.exists(pdf_path) and os.path.getsize(pdf_path) > 1024:
                return True, "libreoffice"
        except Exception:
            pass
    return False, "none"


# ===== 交互 / CLI =====
def prompt(msg: str) -> str:
    val = input(msg).strip()
    if val.upper() == "Q":
        raise KeyboardInterrupt
    if val.upper() == "B":
        return "__BACK__"
    return val


def print_menu() -> None:
    print("—— 交互模式（B=返回，Q=取消）——")
    print("模板固定：", TEMPLATE_PATH_FIXED)
    print("输出目录：", OUT_DIR_FIXED)
    print("默认照片目录：", load_config().get("photos", DEFAULT_PHOTOS_DIR))


def interactive(debug: bool = False) -> None:
    cfg = load_config()
    default_photos_dir = cfg.get("photos", DEFAULT_PHOTOS_DIR)
    cached_photos = expand_photos_arg(default_photos_dir, debug=debug)
    print_menu()

    while True:
        try:
            fields: Dict[str, str] = {}
            for key in PLACEHOLDER_KEYS:
                while True:
                    val = prompt(f"{key}：")
                    if val == "__BACK__":
                        prev_idx = PLACEHOLDER_KEYS.index(key) - 1
                        if prev_idx >= 0:
                            print(f"返回上一项 → 重填 {PLACEHOLDER_KEYS[prev_idx]}")
                        else:
                            print("已在第一项，无法返回。")
                        break
                    if not val:
                        print(f"请输入 {key}。")
                        continue
                    fields[key] = val
                    break
                if key not in fields:
                    break
            if len(fields) < len(PLACEHOLDER_KEYS):
                continue

            hint = "走访照片（输入目录或逗号分隔列表；留空=默认/不插入）："
            photo_input = prompt(hint)
            if photo_input == "__BACK__":
                print("返回上一项 → 重填经营地址")
                continue
            if photo_input.strip():
                photos = expand_photos_arg(photo_input, debug=debug)
                photo_source = photo_input.strip()
            else:
                photos = list(cached_photos)
                photo_source = f"默认 {default_photos_dir}" if cached_photos else "（无照片）"

            print("\n—— 汇总 ——")
            for key in PLACEHOLDER_KEYS:
                print(f"{key}：", fields.get(key, ""))
            print("照片：", len(photos), "张（每页 2 张） 来源：", photo_source)
            go = input("确认生成？[Y/n] ").strip().lower()
            if go and go != "y":
                print("已取消本次生成。")
            else:
                out = generate_report(fields, photos, want_pdf=False, open_out=False, reveal_out=False, debug=debug)
                print(f"✅ 已生成：{out}")
        except KeyboardInterrupt:
            print("\n已退出。")
            return

        cont = input("继续新增？[y/N] ").strip().lower()
        if cont != "y":
            print("已结束。")
            return


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="新开户尽职调查表自动填充")
    parser.add_argument("--customer", help="客户名称")
    parser.add_argument("--industry", help="行业分类")
    parser.add_argument("--address", help="经营地址")
    parser.add_argument("--photos", help="照片目录或逗号分隔文件列表")
    parser.add_argument("--no-photos", action="store_true", help="不插入照片")
    parser.add_argument("--set-photos", help="更新默认照片目录")
    parser.add_argument("--check-template", action="store_true", help="仅检查模板占位符")
    parser.add_argument("--pdf", action="store_true", help="尝试生成 PDF")
    parser.add_argument("--open", action="store_true", help="生成后打开输出目录")
    parser.add_argument("--reveal", action="store_true", help="生成后在 Finder 中定位文件")
    parser.add_argument("--interactive", action="store_true", help="进入交互模式")
    parser.add_argument("--debug", action="store_true", help="打印调试信息")
    return parser.parse_args()


def main() -> None:
    args = parse_args()

    if args.set_photos:
        cfg = load_config()
        cfg["photos"] = args.set_photos
        save_config(cfg)
        print("✅ 默认照片目录已更新：", args.set_photos)
        return

    if args.check_template:
        cfg = load_config()
        template = cfg.get("template", TEMPLATE_PATH_FIXED)
        stats = inspect_template_placeholders(template, PLACEHOLDER_KEYS)
        print("模板占位符统计：")
        for key in PLACEHOLDER_KEYS:
            print(f"  {key}: {stats.get(key, 0)}")
        return

    if len(sys.argv) == 1 or args.interactive:
        interactive(debug=args.debug)
        return

    missing = [
        key
        for key, value in (
            ("客户名称", args.customer),
            ("行业分类", args.industry),
            ("经营地址", args.address),
        )
        if not value
    ]
    if missing:
        raise SystemExit("缺少必填参数：" + ",".join(missing) + "。或使用 --interactive 交互模式。")

    fields = {
        "客户名称": args.customer.strip(),
        "行业分类": args.industry.strip(),
        "经营地址": args.address.strip(),
    }

    if args.no_photos:
        photos: List[str] = []
    else:
        base = args.photos or load_config().get("photos", DEFAULT_PHOTOS_DIR)
        photos = expand_photos_arg(base, debug=args.debug)

    generate_report(
        fields=fields,
        photos=photos,
        want_pdf=args.pdf,
        open_out=args.open,
        reveal_out=args.reveal,
        debug=args.debug,
    )


if __name__ == "__main__":
    main()
